"""
Streamlit frontend for the Statistical Hypothesis Testing Assistant.
"""

import html
import json
import time
import uuid
from io import BytesIO
from typing import Any, Dict, List, Optional

import pandas as pd
import plotly.graph_objects as go
import requests
import streamlit as st
from streamlit_option_menu import option_menu

# ── Configuration ─────────────────────────────────────────────────
import os
BACKEND_URL = os.environ.get("BACKEND_URL", "http://localhost:8000")


def backend(method: str, path: str, **kwargs) -> Optional[Dict]:
    """Make a request to the backend and return JSON or None."""
    try:
        resp = requests.request(method, f"{BACKEND_URL}{path}", timeout=120, **kwargs)
        resp.raise_for_status()
        return resp.json()
    except requests.exceptions.ConnectionError:
        st.error("⚠️ Cannot connect to backend. Is the FastAPI server running on port 8000?")
        return None
    except requests.exceptions.HTTPError as e:
        detail = e.response.text[:300]
        try:
            body = e.response.json()
            if isinstance(body, dict) and "detail" in body:
                detail = body["detail"] if isinstance(body["detail"], str) else str(body["detail"])
        except Exception:
            pass
        st.error(f"Backend error ({e.response.status_code}): {detail}")
        return None
    except Exception as e:
        st.error(f"Unexpected error: {e}")
        return None


# ── Page config ───────────────────────────────────────────────────
st.set_page_config(
    page_title="Statistical Hypothesis Testing Assistant",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Session state initialisation ─────────────────────────────────
for key, default in {
    "pdf_uploaded": False,
    "pdf_filename": "",
    "csv_session_id": None,
    "csv_filename": "",
    "csv_columns": [],
    "csv_dtypes": {},
    "csv_preview": [],
    "csv_rows": 0,
    "chat_history": [],
    "stats_history": [],
    "example_questions": [],
    "stat_question": "",
    "qa_pending_question": None,
    "last_stats_result": None,
    "handbook_concepts": None,
    "handbook_chat_history": [],
}.items():
    if key not in st.session_state:
        st.session_state[key] = default


def _stats_to_bytes(stats_history: list, fmt: str) -> bytes:
    rows = []
    for h in stats_history:
        r = h["result"]
        row = {
            "Question": h["question"],
            "Test": r.get("test_display_name", ""),
            "p-value": r.get("p_value"),
            "Test Statistic": r.get("statistic"),
            "Alpha": r.get("alpha", 0.05),
            "Significant": r.get("significant"),
            "Interpretation": r.get("interpretation", ""),
            "Plain Explanation": r.get("plain_explanation", ""),
            "Rationale": r.get("rationale", ""),
        }
        for role, col in (r.get("variables_used") or {}).items():
            row[f"Variable ({role})"] = col
        rows.append(row)
    df = pd.DataFrame(rows)
    buf = BytesIO()
    if fmt == "csv":
        buf.write(df.to_csv(index=False).encode("utf-8"))
    else:
        with pd.ExcelWriter(buf, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Results")
    buf.seek(0)
    return buf.getvalue()


def _rankings_to_word(question: str, papers: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    buf = BytesIO()
    doc = Document()
    doc.add_heading("Paper Relevance Report", 0)
    doc.add_paragraph(f"Research question: {question}")
    doc.add_paragraph(f"Papers analysed: {len(papers)}")
    doc.add_paragraph()
    for i, p in enumerate(papers, 1):
        doc.add_heading(f"{i}. {p['filename']}", level=2)
        run = doc.add_paragraph().add_run(f"{p['label']}  ·  {int(p['score']*100)}% match")
        run.bold = True
        doc.add_paragraph(p["explanation"])
        if p.get("key_quote") and p["key_quote"] != "—":
            q_para = doc.add_paragraph()
            q_para.add_run(f'"{p["key_quote"]}"').italic = True
        doc.add_paragraph()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _rankings_to_pdf(question: str, papers: list) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    label_colors = {
        "Highly Relevant": colors.HexColor("#0a7c42"),
        "Relevant": colors.HexColor("#1a56a0"),
        "Somewhat Relevant": colors.HexColor("#a05c00"),
        "Less Relevant": colors.HexColor("#c0392b"),
    }
    story = []
    story.append(Paragraph("Paper Relevance Report", styles["Title"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"<b>Research question:</b> {question}", styles["Normal"]))
    story.append(Paragraph(f"<b>Papers analysed:</b> {len(papers)}", styles["Normal"]))
    story.append(Spacer(1, 14))
    for i, p in enumerate(papers, 1):
        story.append(Paragraph(f"{i}. {p['filename']}", styles["Heading2"]))
        col = label_colors.get(p["label"], colors.grey)
        hex_col = "#{:02x}{:02x}{:02x}".format(int(col.red*255), int(col.green*255), int(col.blue*255))
        story.append(Paragraph(
            f'<font color="{hex_col}"><b>{p["label"]}</b></font>  ·  {int(p["score"]*100)}% match',
            styles["Normal"],
        ))
        story.append(Spacer(1, 4))
        story.append(Paragraph(p["explanation"], styles["Normal"]))
        if p.get("key_quote") and p["key_quote"] != "—":
            story.append(Spacer(1, 4))
            story.append(Paragraph(f'<i>"{p["key_quote"]}"</i>', styles["Normal"]))
        story.append(Spacer(1, 8))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
        story.append(Spacer(1, 8))
    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


PDF_QA_PROMPTS = [
    "What is the main research question or objective?",
    "What methodology or study design was used?",
    "Summarize the key findings in a few sentences.",
    "What limitations does the author discuss?",
]


def generate_example_questions(session_id: str, columns: list, dtypes: dict) -> list:
    """
    Fetch dataset-specific example questions from the backend.
    The backend validates each question against the 4 supported tests AND
    the actual data before returning — so every question shown is guaranteed
    to be runnable without an 'unsupported' error.

    Falls back to a local heuristic if the backend call fails.
    """
    result = backend("POST", "/stats/examples", json={"session_id": session_id, "n": 6})
    if result and result.get("questions"):
        return result["questions"]

    # ── Local heuristic fallback ──────────────────────────────────────────
    # Only generates questions for the 4 supported tests using column types.
    numeric_cols = [c for c, d in dtypes.items() if "int" in d or "float" in d]
    binary_cat_cols = []  # we don't have unique counts here, so skip t-test heuristic
    cat_cols = [
        c for c, d in dtypes.items()
        if "int" not in d and "float" not in d and "id" not in c.lower()
    ]

    questions: List[str] = []
    # pearson_correlation
    if len(numeric_cols) >= 2:
        questions.append(f"Is there a linear correlation between {numeric_cols[0]} and {numeric_cols[1]}?")
    # simple_linear_regression
    if len(numeric_cols) >= 2:
        questions.append(f"Does {numeric_cols[0]} predict {numeric_cols[1]}?")
    # chi_square
    if len(cat_cols) >= 2:
        questions.append(f"Is there a significant association between {cat_cols[0]} and {cat_cols[1]}?")
    if len(cat_cols) >= 3:
        questions.append(f"Are {cat_cols[1]} and {cat_cols[2]} statistically independent?")
    # more correlation / regression
    if len(numeric_cols) >= 3:
        questions.append(f"Is there a relationship between {numeric_cols[1]} and {numeric_cols[2]}?")
        questions.append(f"Can {numeric_cols[2]} be predicted from {numeric_cols[0]}?")

    return questions[:6] if questions else ["Ask a statistical question about your dataset."]


def _format_source_excerpt(text: str, max_len: int = 420) -> str:
    text = (text or "").strip()
    if len(text) > max_len:
        text = text[:max_len].rstrip() + "…"
    return html.escape(text)


def _relevance_percent(score: float) -> int:
    return min(100, max(0, int(float(score) * 100)))


def render_source_cards(sources: List[Dict[str, Any]], *, expanded: bool = False) -> None:
    """Render numbered citation cards with relevance bars."""
    if not sources:
        return
    label = "citation" if len(sources) == 1 else "citations"
    with st.expander(f"📎 {len(sources)} document {label}", expanded=expanded):
        for i, src in enumerate(sources, 1):
            pct = _relevance_percent(src.get("score", 0))
            excerpt = _format_source_excerpt(src.get("text", ""))
            st.markdown(
                f"""
                <div class="qa-source-card">
                    <div class="qa-source-header">
                        <span class="qa-citation-badge">{i}</span>
                        <span class="qa-relevance-text">Match</span>
                        <div class="qa-relevance-track" title="Retrieval relevance">
                            <div class="qa-relevance-fill" style="width:{pct}%;"></div>
                        </div>
                        <span class="qa-relevance-pct">{pct}%</span>
                    </div>
                    <p class="qa-source-excerpt">{excerpt}</p>
                </div>
                """,
                unsafe_allow_html=True,
            )


def _render_chart_qa(
    ui_id: str, sid: Optional[str], test_name: str, variables_used: Dict[str, Any],
    chart_key: str, p_val: Optional[float], alpha: float,
) -> None:
    """Mini Q&A thread scoped to a single chart, so users can ask the LLM to explain it further."""
    if not sid:
        return
    chat_key = f"viz_chat_{ui_id}_{chart_key}"
    if chat_key not in st.session_state:
        st.session_state[chat_key] = []

    with st.expander("💬 Ask about this chart"):
        for m in st.session_state[chat_key]:
            with st.chat_message(m["role"]):
                st.markdown(m["content"])

        question = st.text_input(
            "Ask a question about this chart",
            key=f"input_{chat_key}",
            placeholder="e.g. Why is the Treatment box higher than Control?",
            label_visibility="collapsed",
        )
        if st.button("Ask", key=f"btn_{chat_key}") and question.strip():
            history = st.session_state[chat_key][-8:]
            with st.spinner("Thinking…"):
                resp = backend(
                    "POST", "/visualization/ask",
                    json={
                        "session_id": sid,
                        "test_name": test_name,
                        "variables_used": variables_used,
                        "chart_key": chart_key,
                        "question": question,
                        "history": history,
                        "p_value": p_val,
                        "alpha": alpha,
                    },
                )
            st.session_state[chat_key].append({"role": "user", "content": question})
            if resp and resp.get("answer"):
                st.session_state[chat_key].append({"role": "assistant", "content": resp["answer"]})
            st.rerun()


def render_stats_result(result: Dict[str, Any], session_id: Optional[str] = None) -> None:
    """Render a statistical analysis result in a structured, professional layout."""
    test_name = html.escape(str(result.get("test_display_name", "Analysis")))
    rationale = html.escape(str(result.get("rationale", "")))
    interpretation = result.get("interpretation", "—")
    plain = result.get("plain_explanation", "—")

    sig = result.get("significant")
    if sig is True:
        sig_html = '<div class="stats-sig-pill stats-sig-yes">Statistically significant</div>'
    elif sig is False:
        sig_html = '<div class="stats-sig-pill stats-sig-no">Not significant</div>'
    else:
        sig_html = '<div class="stats-sig-pill stats-sig-neutral">Significance N/A</div>'

    p_val = result.get("p_value")
    p_display = f"{p_val:.4f}" if p_val is not None else "—"
    stat_val = result.get("statistic")
    stat_display = f"{stat_val:.4f}" if stat_val is not None else "—"
    alpha = result.get("alpha", 0.05)
    ui_id = result.setdefault("_ui_id", str(uuid.uuid4()))

    with st.container(border=True):
        st.markdown(
            f"""
            <div class="stats-results-header">
                <div class="stats-results-title">
                    <span class="stats-test-badge">{test_name}</span>
                    {sig_html}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        kpi_items = [
            ("p-value", p_display),
            ("Test statistic", stat_display),
            ("α (alpha)", str(alpha)),
        ]
        add = dict(result.get("additional_stats") or {})
        if "n" in add:
            kpi_items.append(("Sample size (n)", str(add["n"])))
        if "r_squared" in add:
            kpi_items.append(("R²", str(add["r_squared"])))
        if "pseudo_r_squared" in add:
            kpi_items.append(("Pseudo R²", str(add["pseudo_r_squared"])))

        kpi_html = '<div class="stats-kpi-row">'
        for label, value in kpi_items[:6]:
            kpi_html += (
                f'<div class="stats-kpi-card">'
                f'<span class="stats-kpi-label">{html.escape(label)}</span>'
                f'<span class="stats-kpi-value">{html.escape(value)}</span>'
                f"</div>"
            )
        kpi_html += "</div>"
        st.markdown(kpi_html, unsafe_allow_html=True)

        sid = session_id or st.session_state.get("csv_session_id")
        if sid:
            suite = backend(
                "POST",
                "/visualization/suite",
                json={
                    "session_id": sid,
                    "test_name": result.get("test_name", ""),
                    "variables_used": result.get("variables_used") or {},
                    "p_value": p_val,
                    "alpha": alpha,
                },
            )
            charts = (suite or {}).get("charts") or []
            if charts:
                st.markdown('<p class="stats-section-label">Visualizations</p>', unsafe_allow_html=True)
                chart_tabs = st.tabs([c["title"] for c in charts])
                for tab, chart in zip(chart_tabs, charts):
                    with tab:
                        st.plotly_chart(
                            go.Figure(chart["figure"]),
                            use_container_width=True,
                            key=f"plot_{ui_id}_{chart['key']}",
                        )
                        if chart.get("interpretation"):
                            st.caption(chart["interpretation"])

                        handbook = chart.get("handbook") or {}
                        if handbook:
                            with st.expander("📖 Learn more about this chart type"):
                                if handbook.get("what"):
                                    st.markdown(f"**What it is:** {handbook['what']}")
                                if handbook.get("how_to_read"):
                                    st.markdown(f"**How to read it:** {handbook['how_to_read']}")
                                if handbook.get("when_used"):
                                    st.markdown(f"**When it's used:** {handbook['when_used']}")

                        _render_chart_qa(
                            ui_id=ui_id,
                            sid=sid,
                            test_name=result.get("test_name", ""),
                            variables_used=result.get("variables_used") or {},
                            chart_key=chart["key"],
                            p_val=p_val,
                            alpha=alpha,
                        )

        if rationale:
            st.markdown(
                f'<div class="stats-rationale-box"><strong>Why this test?</strong> {rationale}</div>',
                unsafe_allow_html=True,
            )

        variables = result.get("variables_used") or {}
        if variables:
            chips = "".join(
                f'<span class="stats-var-chip">'
                f'<span class="stats-var-role">{html.escape(role)}</span>'
                f'{html.escape(str(col))}</span>'
                for role, col in variables.items()
            )
            st.markdown(
                f'<p class="stats-section-label">Variables</p><div class="stats-var-row">{chips}</div>',
                unsafe_allow_html=True,
            )

        if add:
            with st.expander("Detailed statistics", expanded=False):
                add_copy = dict(add)
                if "contingency_table" in add_copy:
                    st.dataframe(pd.DataFrame(add_copy.pop("contingency_table")), use_container_width=True)
                if "coefficients" in add_copy:
                    st.dataframe(pd.DataFrame(add_copy.pop("coefficients")).T, use_container_width=True)
                if "group_stats" in add_copy:
                    st.dataframe(pd.DataFrame(add_copy.pop("group_stats")).T, use_container_width=True)
                rem = {k: v for k, v in add_copy.items() if not isinstance(v, dict)}
                if rem:
                    rcols = st.columns(min(len(rem), 4))
                    for i, (k, v) in enumerate(rem.items()):
                        rcols[i % 4].metric(k.replace("_", " ").title(), str(v))

        checks = result.get("assumption_checks") or []
        if checks:
            with st.expander("Assumption checks", expanded=False):
                for check in checks:
                    passed = check.get("passed", False)
                    cls = "stats-assumption-pass" if passed else "stats-assumption-warn"
                    icon = "✓" if passed else "!"
                    st.markdown(
                        f'<div class="stats-assumption-card {cls}">'
                        f'<span class="stats-assumption-icon">{icon}</span>'
                        f'<div><strong>{html.escape(check.get("name", ""))}</strong><br>'
                        f'<span>{html.escape(check.get("detail", ""))}</span></div>'
                        f"</div>",
                        unsafe_allow_html=True,
                    )

        st.markdown('<p class="stats-section-label">Interpretation</p>', unsafe_allow_html=True)
        icol1, icol2 = st.columns(2, gap="medium")
        with icol1:
            st.markdown("##### Technical")
            with st.container(border=True):
                st.markdown(interpretation)
        with icol2:
            st.markdown("##### Plain language")
            with st.container(border=True):
                st.markdown(plain)


def handle_qa_question(question: str) -> None:
    """Send a question to the PDF Q&A API and append results to chat history."""
    question = question.strip()
    if not question or not st.session_state.pdf_uploaded:
        return

    st.session_state.chat_history.append({"role": "user", "content": question})
    result = backend("POST", "/qa/ask", json={"question": question, "top_k": 5})
    if result:
        st.session_state.chat_history.append({
            "role": "assistant",
            "content": result.get("answer", "No answer returned."),
            "sources": result.get("sources", []),
        })


# ── Custom CSS ────────────────────────────────────────────────────
st.markdown("""
<style>
    :root {
        --bg-app: #10131c;
        --bg-sidebar: #0b0e15;
        --bg-card: #161a24;
        --bg-card-alt: #1b2130;
        --border: #242a38;
        --border-strong: #323a4d;
        --accent: #4f7cff;
        --accent-hover: #6b91ff;
        --accent-soft: rgba(79, 124, 255, 0.14);
        --text-primary: #e7e9f0;
        --text-secondary: #99a2b8;
        --text-muted: #6b7488;
        --success-bg: rgba(52, 211, 153, 0.12);
        --success-text: #34d399;
        --success-border: rgba(52, 211, 153, 0.35);
        --danger-bg: rgba(248, 113, 113, 0.12);
        --danger-text: #f87171;
        --danger-border: rgba(248, 113, 113, 0.35);
        --warning-bg: rgba(251, 191, 36, 0.12);
        --warning-text: #fbbf24;
        --warning-border: rgba(251, 191, 36, 0.35);
    }

    .stApp { background: var(--bg-app); }
    .block-container { padding-top: 3.5rem; max-width: 1180px; }

    /* ── Sidebar shell ── */
    div[data-testid="stSidebar"] {
        background: var(--bg-sidebar);
        border-right: 1px solid var(--border);
    }
    div[data-testid="stSidebar"] > div { padding-top: 1.1rem; }

    .sidebar-brand {
        display: flex;
        align-items: center;
        gap: 0.6rem;
        padding: 0 0.5rem 1.1rem;
        margin-bottom: 0.6rem;
        border-bottom: 1px solid var(--border);
    }
    .sidebar-brand-icon {
        width: 34px;
        height: 34px;
        border-radius: 9px;
        background: var(--accent-soft);
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 1.05rem;
        flex-shrink: 0;
    }
    .sidebar-brand-text {
        display: flex;
        flex-direction: column;
        line-height: 1.2;
    }
    .sidebar-brand-text span:first-child {
        font-weight: 700;
        font-size: 0.95rem;
        color: var(--text-primary);
    }
    .sidebar-brand-text span:last-child {
        font-size: 0.78rem;
        color: var(--text-secondary);
    }

    .sidebar-section-label {
        color: var(--text-muted);
        font-size: 0.72rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        margin: 1.2rem 0 0.5rem 0.25rem;
    }
    .sidebar-upload-label {
        color: var(--text-secondary);
        font-size: 0.82rem;
        margin: 0 0 0.3rem 0.1rem;
    }
    .sidebar-upload-label em { color: var(--text-muted); font-style: normal; }

    div[data-testid="stSidebar"] div[data-testid="stFileUploaderDropzone"] {
        background: var(--bg-card);
        border: 1px dashed var(--border-strong);
        border-radius: 10px;
        padding: 0.5rem 0.75rem;
        min-height: 0;
    }
    div[data-testid="stSidebar"] div[data-testid="stFileUploaderDropzoneInstructions"] span,
    div[data-testid="stSidebar"] div[data-testid="stFileUploaderDropzoneInstructions"] small {
        font-size: 0.72rem;
        color: var(--text-muted);
    }
    div[data-testid="stSidebar"] section[data-testid="stFileUploaderDropzone"] svg { display: none; }

    .sidebar-footer {
        display: flex;
        align-items: center;
        gap: 0.6rem;
        margin-top: 1.4rem;
        padding: 0.85rem 0.5rem 0.2rem;
        border-top: 1px solid var(--border);
    }
    .sidebar-avatar {
        width: 32px;
        height: 32px;
        border-radius: 50%;
        background: var(--accent);
        color: #fff;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: 700;
        font-size: 0.85rem;
        flex-shrink: 0;
    }
    .sidebar-user-meta { display: flex; flex-direction: column; line-height: 1.25; min-width: 0; }
    .sidebar-user-meta span:first-child { font-size: 0.85rem; color: var(--text-primary); font-weight: 600; }
    .sidebar-user-meta span:last-child {
        font-size: 0.74rem;
        color: var(--text-muted);
        overflow: hidden;
        text-overflow: ellipsis;
        white-space: nowrap;
    }

    /* ── Top bar ── */
    .topbar-icons {
        text-align: right;
        font-size: 1.15rem;
        padding-top: 0.35rem;
        opacity: 0.85;
    }
    div[data-testid="stTextInput"] input {
        background: var(--bg-card) !important;
        border: 1px solid var(--border-strong) !important;
        border-radius: 999px !important;
        color: var(--text-primary) !important;
        padding: 0.55rem 1.1rem !important;
    }
    div[data-testid="stTextInput"] input::placeholder { color: var(--text-muted) !important; }

    /* ── Hero banners (page titles) ── */
    .qa-hero, .stats-hero {
        background: linear-gradient(135deg, #1a2340 0%, #24305a 100%);
        border: 1px solid var(--border-strong);
        border-radius: 14px;
        padding: 1.25rem 1.5rem;
        margin-bottom: 1rem;
        color: var(--text-primary);
    }
    .qa-hero h3, .stats-hero h3 { margin: 0 0 0.35rem; font-size: 1.35rem; font-weight: 600; color: #fff; }
    .qa-hero p, .stats-hero p { margin: 0; font-size: 0.92rem; opacity: 0.85; line-height: 1.5; }
    .qa-doc-pill, .stats-dataset-pill {
        display: inline-block;
        margin-top: 0.75rem;
        background: rgba(255,255,255,0.08);
        border: 1px solid rgba(255,255,255,0.18);
        padding: 0.35rem 0.85rem;
        border-radius: 999px;
        font-size: 0.82rem;
        color: #dbe2f5;
    }

    [data-testid="stVerticalBlockBorderWrapper"] {
        border-color: var(--border) !important;
        border-radius: 12px !important;
        background: var(--bg-card);
    }
    [data-testid="stExpander"] {
        border: 1px solid var(--border) !important;
        border-radius: 10px !important;
        background: var(--bg-card);
    }

    .qa-toolbar-meta { color: var(--text-secondary); font-size: 0.85rem; margin: 0; padding-top: 0.35rem; }
    .qa-empty-state, .stats-empty-state {
        text-align: center;
        padding: 2.5rem 1.5rem;
        background: var(--bg-card);
        border: 1px dashed var(--border-strong);
        border-radius: 12px;
        margin: 0.5rem 0 1rem;
    }
    .qa-empty-state h4, .stats-empty-state h4 { color: var(--text-primary); margin: 0 0 0.5rem; font-size: 1.1rem; }
    .qa-empty-state p, .stats-empty-state p {
        color: var(--text-secondary);
        margin: 0 0 1rem;
        font-size: 0.9rem;
        max-width: 32rem;
        margin-left: auto;
        margin-right: auto;
        line-height: 1.55;
    }
    .qa-steps, .stats-empty-state ol {
        display: inline-block;
        text-align: left;
        color: var(--text-secondary);
        font-size: 0.88rem;
        line-height: 1.8;
        margin: 0;
        padding-left: 1.2rem;
    }
    .qa-welcome {
        background: var(--bg-card);
        border-left: 4px solid var(--accent);
        border-radius: 0 10px 10px 0;
        padding: 1rem 1.15rem;
        margin-bottom: 1rem;
        color: var(--text-secondary);
        font-size: 0.9rem;
        line-height: 1.55;
    }
    .qa-welcome strong { color: var(--text-primary); }
    .qa-prompt-label, .stats-section-label {
        color: var(--text-muted);
        font-size: 0.78rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.07em;
        margin: 0.75rem 0 0.5rem;
    }
    .qa-source-card {
        background: var(--bg-card-alt);
        border: 1px solid var(--border);
        border-left: 4px solid var(--accent);
        border-radius: 8px;
        padding: 0.85rem 1rem;
        margin-bottom: 0.65rem;
    }
    .qa-source-header { display: flex; align-items: center; gap: 0.5rem; margin-bottom: 0.5rem; flex-wrap: wrap; }
    .qa-citation-badge {
        background: var(--accent);
        color: #fff;
        font-size: 0.72rem;
        font-weight: 700;
        min-width: 1.5rem;
        height: 1.5rem;
        line-height: 1.5rem;
        text-align: center;
        border-radius: 6px;
    }
    .qa-relevance-text { font-size: 0.72rem; color: var(--text-secondary); font-weight: 600; text-transform: uppercase; letter-spacing: 0.04em; }
    .qa-relevance-track { flex: 1; min-width: 80px; max-width: 140px; height: 6px; background: var(--border); border-radius: 999px; overflow: hidden; }
    .qa-relevance-fill { height: 100%; background: linear-gradient(90deg, var(--accent), var(--accent-hover)); border-radius: 999px; }
    .qa-relevance-pct { font-size: 0.75rem; color: var(--accent-hover); font-weight: 600; min-width: 2.2rem; }
    .qa-source-excerpt { margin: 0; font-size: 0.88rem; line-height: 1.6; color: var(--text-secondary); }
    div[data-testid="stChatMessage"] {
        background: var(--bg-card);
        border: 1px solid var(--border);
        border-radius: 12px;
    }

    /* ── Statistical Analysis ── */
    .stats-input-hint { color: var(--text-secondary); font-size: 0.88rem; margin: 0 0 0.75rem; line-height: 1.5; }
    .stats-results-header { margin-bottom: 1rem; }
    .stats-results-title { display: flex; flex-wrap: wrap; align-items: center; gap: 0.65rem; }
    .stats-test-badge { background: var(--accent); color: #fff; padding: 0.4rem 0.95rem; border-radius: 8px; font-size: 0.95rem; font-weight: 600; }
    .stats-sig-pill { padding: 0.35rem 0.85rem; border-radius: 999px; font-size: 0.8rem; font-weight: 700; letter-spacing: 0.02em; }
    .stats-sig-yes { background: var(--success-bg); color: var(--success-text); border: 1px solid var(--success-border); }
    .stats-sig-no { background: var(--danger-bg); color: var(--danger-text); border: 1px solid var(--danger-border); }
    .stats-sig-neutral { background: var(--bg-card-alt); color: var(--text-secondary); border: 1px solid var(--border); }
    .stats-kpi-row { display: grid; grid-template-columns: repeat(auto-fit, minmax(120px, 1fr)); gap: 0.65rem; margin-bottom: 1rem; }
    .stats-kpi-card { background: var(--bg-card-alt); border: 1px solid var(--border); border-radius: 10px; padding: 0.75rem 0.9rem; border-left: 4px solid var(--accent); }
    .stats-kpi-label { display: block; font-size: 0.72rem; color: var(--text-secondary); text-transform: uppercase; letter-spacing: 0.05em; font-weight: 600; margin-bottom: 0.25rem; }
    .stats-kpi-value { display: block; font-size: 1.25rem; font-weight: 700; color: var(--text-primary); }
    .stats-rationale-box {
        background: var(--bg-card-alt);
        border: 1px solid var(--border);
        border-left: 4px solid var(--accent);
        border-radius: 0 10px 10px 0;
        padding: 0.9rem 1.1rem;
        color: var(--text-secondary);
        font-size: 0.9rem;
        line-height: 1.55;
        margin-bottom: 0.5rem;
    }
    .stats-rationale-box strong { color: var(--text-primary); }
    .stats-var-row { display: flex; flex-wrap: wrap; gap: 0.5rem; margin-bottom: 0.5rem; }
    .stats-var-chip { background: var(--bg-card-alt); border: 1px solid var(--border-strong); border-radius: 8px; padding: 0.35rem 0.7rem; font-size: 0.85rem; color: var(--text-primary); }
    .stats-var-role { color: var(--text-muted); font-size: 0.72rem; text-transform: uppercase; font-weight: 600; margin-right: 0.35rem; }
    .stats-assumption-card { display: flex; gap: 0.65rem; align-items: flex-start; padding: 0.7rem 0.9rem; border-radius: 8px; margin-bottom: 0.5rem; font-size: 0.88rem; line-height: 1.45; }
    .stats-assumption-pass { background: var(--success-bg); border: 1px solid var(--success-border); color: #a9e8cd; }
    .stats-assumption-warn { background: var(--warning-bg); border: 1px solid var(--warning-border); color: #f3d78c; }
    .stats-assumption-icon { font-weight: 700; width: 1.25rem; height: 1.25rem; border-radius: 50%; display: flex; align-items: center; justify-content: center; flex-shrink: 0; background: rgba(255,255,255,0.08); }
    .stats-unsupported-box {
        background: var(--warning-bg); border: 1px solid var(--warning-border);
        border-left: 4px solid var(--warning-text); border-radius: 0 10px 10px 0;
        padding: 1rem 1.15rem; color: #f3d78c; line-height: 1.55; margin-top: 1rem;
    }
    .stats-error-box {
        background: var(--danger-bg); border: 1px solid var(--danger-border);
        border-left: 4px solid var(--danger-text); border-radius: 0 10px 10px 0;
        padding: 1rem 1.15rem; color: #f5b3b3; line-height: 1.55; margin-top: 1rem;
    }
    .stats-history-item { color: var(--text-secondary); font-size: 0.85rem; margin-bottom: 0.35rem; }
    .stats-history-item strong { color: var(--text-primary); }
    .stats-example-row [data-testid="column"] button {
        min-height: 3.1rem; white-space: normal !important; line-height: 1.35 !important; text-align: left !important;
    }
    [data-testid="stVerticalBlock"]:has([data-testid="stSelectbox"]) [data-testid="stHorizontalBlock"] { align-items: flex-end; }
    [data-testid="stVerticalBlock"]:has([data-testid="stSelectbox"]) [data-testid="stHorizontalBlock"] button { margin-top: 0; min-height: 2.5rem; }

    /* ── Statistics Handbook ── */
    .handbook-ask-card {
        background: var(--bg-card);
        border: 1px solid var(--border);
        border-radius: 14px;
        padding: 1.1rem 1.25rem;
        margin-bottom: 1.25rem;
    }
    .handbook-ask-title { display: flex; align-items: center; gap: 0.5rem; font-weight: 700; font-size: 1rem; color: var(--text-primary); margin-bottom: 0.15rem; }
    .handbook-ask-subtitle { color: var(--text-secondary); font-size: 0.85rem; margin-bottom: 0.9rem; }
    .handbook-answer-box {
        background: var(--bg-card-alt);
        border: 1px solid var(--border);
        border-radius: 10px;
        padding: 1rem 1.15rem;
        margin-top: 0.9rem;
        color: var(--text-secondary);
        font-size: 0.92rem;
        line-height: 1.6;
    }
    .concept-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(280px, 1fr)); gap: 0.85rem; margin-bottom: 1.5rem; }
    .concept-card { background: var(--bg-card); border: 1px solid var(--border); border-radius: 12px; padding: 1rem 1.15rem; }
    .concept-card-title { display: flex; align-items: center; gap: 0.5rem; font-weight: 700; font-size: 0.92rem; color: var(--text-primary); margin-bottom: 0.4rem; }
    .concept-card-icon {
        width: 26px; height: 26px; border-radius: 7px; background: var(--accent-soft);
        display: flex; align-items: center; justify-content: center; font-size: 0.85rem; flex-shrink: 0;
    }
    .concept-card-body { color: var(--text-secondary); font-size: 0.86rem; line-height: 1.55; }

    /* ── Dataset result cards (Find Papers & Datasets) ── */
    .dataset-format-chip {
        background: var(--bg-card-alt); color: var(--text-secondary); border: 1px solid var(--border-strong);
        font-size: 0.7rem; font-weight: 600; padding: 0.15rem 0.5rem; border-radius: 999px; margin-right: 0.3rem;
    }
</style>
""", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════════
# SIDEBAR — Navigation + Uploads
# ════════════════════════════════════════════════════════════════════

NAV_OPTIONS = [
    "Hypothesis Chat",
    "Ask Questions on Research Papers",
    "Find Papers & Datasets",
    "Rank Papers",
    "Statistical Analysis",
    "Statistics Handbook",
    "Data Preview",
]
NAV_ICONS = [
    "chat-square-text",
    "chat-dots",
    "search",
    "bar-chart-line",
    "clipboard-data",
    "book",
    "table",
]

if "nav_selection" not in st.session_state:
    st.session_state.nav_selection = NAV_OPTIONS[0]

with st.sidebar:
    st.markdown("""
    <div class="sidebar-brand">
        <div class="sidebar-brand-icon">📊</div>
        <div class="sidebar-brand-text">
            <span>Hypothesis</span>
            <span>Testing Assistant</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

    manual_idx = None
    if st.session_state.pop("_force_nav", False):
        manual_idx = NAV_OPTIONS.index(st.session_state.nav_selection)

    selected_page = option_menu(
        menu_title=None,
        options=NAV_OPTIONS,
        icons=NAV_ICONS,
        default_index=NAV_OPTIONS.index(st.session_state.nav_selection),
        manual_select=manual_idx,
        styles={
            "container": {"padding": "0", "background-color": "transparent"},
            "icon": {"color": "#8b93a7", "font-size": "0.95rem"},
            "nav-link": {
                "font-size": "0.9rem",
                "color": "#c7cddb",
                "text-align": "left",
                "margin": "0.15rem 0",
                "padding": "0.55rem 0.75rem",
                "border-radius": "8px",
                "--hover-color": "#1a2030",
            },
            "nav-link-selected": {
                "background-color": "rgba(79,124,255,0.16)",
                "color": "#8fb0ff",
                "font-weight": "600",
            },
        },
        key="nav_menu",
    )
    st.session_state.nav_selection = selected_page

    st.markdown('<p class="sidebar-section-label">Upload</p>', unsafe_allow_html=True)

    # ── PDF (optional) ───────────────────────────────────────────
    st.markdown(
        '<p class="sidebar-upload-label">📄 Research PDF <em>(optional · max 200MB)</em></p>',
        unsafe_allow_html=True,
    )
    pdf_file = st.file_uploader(
        "Upload a strategy/research PDF",
        type=["pdf"],
        key="pdf_uploader",
        help="The PDF will be indexed for Q&A. Useful for methodology papers.",
        label_visibility="collapsed",
    )

    if pdf_file and not st.session_state.pdf_uploaded:
        with st.spinner("Parsing and indexing PDF…"):
            result = backend(
                "POST", "/upload/pdf",
                files={"file": (pdf_file.name, pdf_file.getvalue(), "application/pdf")},
            )
        if result and result.get("success"):
            st.session_state.pdf_uploaded = True
            st.session_state.pdf_filename = pdf_file.name
            st.success(f"✅ {result.get('chunks_created', 0)} chunks indexed")
        elif result:
            st.error(result.get("message", "Upload failed"))

    if st.session_state.pdf_uploaded:
        st.markdown(f"✅ **{st.session_state.pdf_filename}** indexed")
        if st.button("Remove PDF"):
            st.session_state.pdf_uploaded = False
            st.session_state.pdf_filename = ""
            st.rerun()

    # ── CSV / XLSX ────────────────────────────────────────────────
    st.markdown(
        '<p class="sidebar-upload-label" style="margin-top:0.9rem;">📊 Dataset <em>(CSV / XLSX)</em></p>',
        unsafe_allow_html=True,
    )
    csv_file = st.file_uploader(
        "Upload your data file",
        type=["csv", "xlsx", "xls"],
        key="csv_uploader",
        label_visibility="collapsed",
    )

    if csv_file:
        fname = csv_file.name
        if fname != st.session_state.csv_filename:
            with st.spinner("Loading dataset…"):
                result = backend(
                    "POST", "/upload/csv",
                    files={"file": (fname, csv_file.getvalue(), "text/csv")},
                )
            if result and result.get("success"):
                st.session_state.csv_session_id = result["session_id"]
                st.session_state.csv_filename = fname
                st.session_state.csv_columns = result["columns"]
                st.session_state.csv_dtypes = result["dtypes"]
                st.session_state.csv_preview = result["preview"]
                st.session_state.csv_rows = result["rows"]
                # Reset examples so they get regenerated for the new dataset
                st.session_state.example_questions = []
                st.session_state.stat_question = ""
                st.success(f"✅ {result['rows']} rows × {len(result['columns'])} cols")
            elif result:
                st.error(result.get("message", "Upload failed"))

    if st.session_state.csv_session_id:
        st.markdown(f"✅ **{st.session_state.csv_filename}**")
        st.caption(f"{st.session_state.csv_rows} rows · {len(st.session_state.csv_columns)} columns")
        with st.expander("Column info"):
            for col in st.session_state.csv_columns:
                dtype = st.session_state.csv_dtypes.get(col, "")
                icon = "🔢" if "int" in dtype or "float" in dtype else "🔤"
                st.caption(f"{icon} `{col}` — {dtype}")

        if st.button("Remove dataset"):
            st.session_state.csv_session_id = None
            st.session_state.csv_filename = ""
            st.session_state.csv_columns = []
            st.session_state.csv_dtypes = {}
            st.session_state.csv_preview = []
            st.session_state.csv_rows = 0
            st.session_state.example_questions = []
            st.session_state.stat_question = ""
            st.rerun()

    st.markdown("""
    <div class="sidebar-footer">
        <div class="sidebar-avatar">C</div>
        <div class="sidebar-user-meta">
            <span>Chiara von Watzdorf</span>
            <span>chiara.vonwatzdorf@gmail.com</span>
        </div>
    </div>
    """, unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════════
# MAIN AREA — Top bar + page dispatch
# ════════════════════════════════════════════════════════════════════

top_col_search, top_col_icons = st.columns([6, 1], vertical_alignment="center")
with top_col_search:
    global_search = st.text_input(
        "Global search",
        placeholder="🔍  Search papers, datasets, or ask a question…",
        key="global_search",
        label_visibility="collapsed",
    )
with top_col_icons:
    st.markdown('<div class="topbar-icons">🔔&nbsp;&nbsp;⚙️</div>', unsafe_allow_html=True)

if global_search and global_search != st.session_state.get("_last_global_search", ""):
    st.session_state["_last_global_search"] = global_search
    st.session_state["find_hypothesis"] = global_search
    st.session_state["find_mode"] = "Research Papers"
    st.session_state["nav_selection"] = "Find Papers & Datasets"
    st.session_state["_force_nav"] = True
    st.rerun()


# ── Tab 1: Hypothesis Chatbot ────────────────────────────────────
if selected_page == "Hypothesis Chat":
    st.markdown("""
    <div class="qa-hero">
        <h3>Hypothesis Chat</h3>
        <p>Throw any hypothesis at me — scientific, philosophical, or completely wild.
        I'll engage with it seriously, creatively, and honestly.</p>
    </div>
    """, unsafe_allow_html=True)

    if "hyp_chat_history" not in st.session_state:
        st.session_state.hyp_chat_history = []

    starter_hypotheses = [
        "Will there be unicorns one day?",
        "Could humans ever photosynthesize like plants?",
        "Is consciousness just an illusion?",
        "Will AI become smarter than all humans combined?",
    ]

    if not st.session_state.hyp_chat_history:
        st.markdown('<p class="qa-prompt-label">Try one of these</p>', unsafe_allow_html=True)
        s_cols = st.columns(2, gap="small")
        for idx, s in enumerate(starter_hypotheses):
            with s_cols[idx % 2]:
                if st.button(s, key=f"starter_{idx}", use_container_width=True, type="secondary"):
                    st.session_state.hyp_chat_history.append({"role": "user", "content": s})
                    with st.spinner("Thinking…"):
                        result = backend("POST", "/chat/message", json={
                            "message": s,
                            "history": [],
                        })
                    if result:
                        st.session_state.hyp_chat_history.append({
                            "role": "assistant",
                            "content": result["reply"],
                        })
                    st.rerun()

    for msg in st.session_state.hyp_chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    if st.session_state.hyp_chat_history:
        if st.button("Clear chat", key="clear_hyp_chat"):
            st.session_state.hyp_chat_history = []
            st.rerun()

    user_input = st.chat_input("Type any hypothesis, however wild…", key="hyp_chat_input")
    if user_input:
        st.session_state.hyp_chat_history.append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.markdown(user_input)
        with st.chat_message("assistant"):
            with st.spinner("Thinking…"):
                result = backend("POST", "/chat/message", json={
                    "message": user_input,
                    "history": st.session_state.hyp_chat_history[:-1],
                })
            if result:
                reply = result["reply"]
                st.markdown(reply)
                st.session_state.hyp_chat_history.append({"role": "assistant", "content": reply})


# ── Tab 2: Ask Questions on Research Papers ──────────────────────
elif selected_page == "Ask Questions on Research Papers":
    pdf_ready = st.session_state.pdf_uploaded
    pdf_name = st.session_state.pdf_filename or "Document"
    msg_count = len(st.session_state.chat_history)

    # Hero header
    hero_doc = (
        f'<span class="qa-doc-pill">📄 {html.escape(pdf_name)}</span>'
        if pdf_ready
        else ""
    )
    st.markdown(
        f"""
        <div class="qa-hero">
            <h3>Document Q&amp;A</h3>
            <p>Ask natural-language questions about your research PDF.
            Answers are grounded in retrieved passages from the indexed document.</p>
            {hero_doc}
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Toolbar
    t_left, t_right = st.columns([3, 1], vertical_alignment="center")
    with t_left:
        if pdf_ready:
            st.markdown(
                f'<p class="qa-toolbar-meta">'
                f'{"●" if msg_count else "○"} '
                f'{msg_count} message{"s" if msg_count != 1 else ""} in this session'
                f'</p>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<p class="qa-toolbar-meta">No document indexed — upload a PDF in the sidebar</p>',
                unsafe_allow_html=True,
            )
    with t_right:
        if pdf_ready and st.session_state.chat_history:
            if st.button("Clear chat", key="clear_chat", use_container_width=True):
                st.session_state.chat_history = []
                st.session_state.qa_pending_question = None
                st.rerun()

    if not pdf_ready:
        st.markdown(
            """
            <div class="qa-empty-state">
                <h4>No document loaded yet</h4>
                <p>Upload a research PDF in the sidebar to index it for semantic search and chat.</p>
                <ol class="qa-steps">
                    <li>Open the sidebar and choose <strong>Research PDF</strong></li>
                    <li>Wait for parsing and indexing to complete</li>
                    <li>Return here to ask questions about the paper</li>
                </ol>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        with st.container(border=True):
            # Welcome + starter prompts when chat is empty
            if not st.session_state.chat_history:
                st.markdown(
                    f"""
                    <div class="qa-welcome">
                        <strong>{html.escape(pdf_name)}</strong> is ready.
                        Ask anything about methods, findings, definitions, or limitations —
                        or start with a suggested prompt below.
                    </div>
                    <p class="qa-prompt-label">Suggested questions</p>
                    """,
                    unsafe_allow_html=True,
                )
                pcols = st.columns(2, gap="small", vertical_alignment="top")
                for idx, prompt in enumerate(PDF_QA_PROMPTS):
                    with pcols[idx % 2]:
                        if st.button(
                            prompt,
                            key=f"qa_prompt_{idx}",
                            use_container_width=True,
                            type="secondary",
                        ):
                            st.session_state.qa_pending_question = prompt
                            st.rerun()

            # Process pending question (from suggested prompts)
            pending = st.session_state.qa_pending_question
            if pending:
                st.session_state.qa_pending_question = None
                with st.spinner("Searching document and generating answer…"):
                    handle_qa_question(pending)
                st.rerun()

            # Chat history
            for msg in st.session_state.chat_history:
                with st.chat_message(msg["role"]):
                    st.markdown(msg["content"])
                    if msg["role"] == "assistant" and msg.get("sources"):
                        render_source_cards(msg["sources"])

    question = st.chat_input(
        "Ask about methodology, findings, limitations…",
        disabled=not pdf_ready,
        key="qa_chat_input",
    )

    if question:
        with st.spinner("Searching document and generating answer…"):
            handle_qa_question(question)
        st.rerun()


# ── Tab 3: Find Papers & Datasets ────────────────────────────────
elif selected_page == "Find Papers & Datasets":
    st.markdown("""
    <div class="qa-hero">
        <h3>Find Papers &amp; Datasets</h3>
        <p>Enter your hypothesis or research question — the AI extracts keywords and
        searches academic literature or open dataset registries for the most relevant matches.</p>
    </div>
    """, unsafe_allow_html=True)

    for key, default in {
        "search_results": None,
        "search_query_used": "",
        "search_summary": "",
        "dataset_results": None,
        "dataset_query_used": "",
        "dataset_summary": "",
    }.items():
        if key not in st.session_state:
            st.session_state[key] = default

    find_mode = st.radio(
        "What are you looking for?",
        ["Research Papers", "Datasets"],
        horizontal=True,
        key="find_mode",
    )

    with st.container(border=True):
        hypothesis = st.text_area(
            "Your hypothesis or research question",
            placeholder="e.g. Does sleep deprivation impair working memory in young adults?",
            height=100,
            label_visibility="collapsed",
            key="find_hypothesis",
        )
        s_col1, s_col2 = st.columns([2, 1], gap="medium", vertical_alignment="bottom")
        with s_col1:
            result_limit = st.selectbox("Results to return", [5, 10, 15, 20], index=1)
        with s_col2:
            search_btn = st.button(
                "Search Papers" if find_mode == "Research Papers" else "Search Datasets",
                type="primary",
                disabled=not hypothesis.strip(),
                use_container_width=True,
            )

        st.markdown('<p class="stats-section-label" style="margin-top:0.75rem;">Sources</p>', unsafe_allow_html=True)
        src_cols = st.columns(4)
        if find_mode == "Research Papers":
            all_sources = ["Semantic Scholar", "arXiv", "PubMed", "OpenAlex"]
        else:
            all_sources = ["Zenodo", "DataCite", "OpenAlex Datasets"]
        selected_sources = []
        for i, src in enumerate(all_sources):
            with src_cols[i]:
                if st.checkbox(src, value=True, key=f"src_{find_mode}_{src}"):
                    selected_sources.append(src)

    if search_btn and hypothesis.strip():
        if not selected_sources:
            st.warning("Please select at least one source.")
        elif find_mode == "Research Papers":
            with st.spinner("Searching across selected databases…"):
                resp = backend("POST", "/search/papers", json={
                    "question": hypothesis,
                    "limit": result_limit,
                    "sources": selected_sources,
                })
                if resp:
                    st.session_state.search_results = resp.get("results", [])
                    st.session_state.search_query_used = resp.get("query_used", "")
                    st.session_state.search_summary = resp.get("summary", "")
        else:
            with st.spinner("Searching open dataset registries…"):
                resp = backend("POST", "/search/datasets", json={
                    "question": hypothesis,
                    "limit": result_limit,
                    "sources": selected_sources,
                })
                if resp:
                    st.session_state.dataset_results = resp.get("results", [])
                    st.session_state.dataset_query_used = resp.get("query_used", "")
                    st.session_state.dataset_summary = resp.get("summary", "")

    if find_mode == "Research Papers" and st.session_state.search_results is not None:
        query_used = st.session_state.search_query_used
        results = st.session_state.search_results

        st.markdown(
            f'<p class="qa-toolbar-meta">Search query: <strong>{html.escape(query_used)}</strong> · '
            f'{len(results)} result{"s" if len(results) != 1 else ""} returned</p>',
            unsafe_allow_html=True,
        )

        if st.session_state.search_summary:
            st.markdown(
                f"""
                <div class="stats-rationale-box" style="margin:0.75rem 0 1rem;">
                    <strong>Literature summary</strong><br><br>
                    {html.escape(st.session_state.search_summary)}
                </div>
                """,
                unsafe_allow_html=True,
            )

        if not results:
            st.info("No papers found. Try rephrasing your hypothesis.")
        else:
            for i, paper in enumerate(results, 1):
                title = html.escape(paper.get("title") or "Untitled")
                authors = paper.get("authors") or []
                author_str = html.escape(", ".join(authors[:3]) + (" et al." if len(authors) > 3 else ""))
                year = paper.get("year") or "—"
                citations = paper.get("citation_count")
                citation_str = f"{citations:,}" if citations is not None else "—"
                abstract = (paper.get("abstract") or "").strip()
                if len(abstract) > 400:
                    abstract = abstract[:400].rstrip() + "…"
                abstract_html = html.escape(abstract) if abstract else "<em>No abstract available.</em>"
                doi = paper.get("doi")
                pdf_url = paper.get("pdf_url")
                paper_url = paper.get("paper_url")

                link_parts = []
                if paper_url:
                    link_parts.append(f'<a href="{html.escape(paper_url)}" target="_blank">View on Semantic Scholar</a>')
                if pdf_url:
                    link_parts.append(f'<a href="{html.escape(pdf_url)}" target="_blank">Open PDF</a>')
                if doi:
                    link_parts.append(f'<a href="https://doi.org/{html.escape(doi)}" target="_blank">DOI</a>')
                links_html = " · ".join(link_parts) if link_parts else ""

                source = html.escape(paper.get("source") or "")
                source_colors = {
                    "Semantic Scholar": "#2d6a9f",
                    "arXiv": "#b31b1b",
                    "PubMed": "#2e7d32",
                    "OpenAlex": "#6a1b9a",
                }
                source_color = source_colors.get(paper.get("source", ""), "#555")
                source_badge = (
                    f'<span style="background:{source_color};color:#fff;font-size:0.7rem;'
                    f'font-weight:700;padding:0.2rem 0.55rem;border-radius:999px;margin-left:0.5rem;">'
                    f'{source}</span>'
                ) if source else ""

                with st.expander(f"{i}. {paper.get('title', 'Untitled')} ({year})", expanded=(i == 1)):
                    st.markdown(
                        f"""
                        <div style="margin-bottom:0.5rem;">
                            {source_badge}
                            <span style="color:var(--text-secondary);font-size:0.85rem;margin-left:0.35rem;">{author_str}</span>
                            &nbsp;·&nbsp;
                            <span style="color:var(--text-secondary);font-size:0.85rem;">{year}</span>
                            &nbsp;·&nbsp;
                            <span style="color:var(--text-secondary);font-size:0.85rem;">🔖 {citation_str} citations</span>
                        </div>
                        <p style="font-size:0.9rem;line-height:1.6;color:var(--text-secondary);margin-bottom:0.75rem;">{abstract_html}</p>
                        <p style="font-size:0.82rem;">{links_html}</p>
                        """,
                        unsafe_allow_html=True,
                    )

        if st.button("Clear results", key="clear_search"):
            st.session_state.search_results = None
            st.session_state.search_query_used = ""
            st.session_state.search_summary = ""
            st.rerun()

    if find_mode == "Datasets" and st.session_state.dataset_results is not None:
        query_used = st.session_state.dataset_query_used
        results = st.session_state.dataset_results

        st.markdown(
            f'<p class="qa-toolbar-meta">Search query: <strong>{html.escape(query_used)}</strong> · '
            f'{len(results)} result{"s" if len(results) != 1 else ""} returned</p>',
            unsafe_allow_html=True,
        )

        if st.session_state.dataset_summary:
            st.markdown(
                f"""
                <div class="stats-rationale-box" style="margin:0.75rem 0 1rem;">
                    <strong>Dataset fit summary</strong><br><br>
                    {html.escape(st.session_state.dataset_summary)}
                </div>
                """,
                unsafe_allow_html=True,
            )

        if not results:
            st.info("No datasets found. Try rephrasing your hypothesis.")
        else:
            for i, ds in enumerate(results, 1):
                authors = ds.get("authors") or []
                author_str = html.escape(", ".join(authors[:3]) + (" et al." if len(authors) > 3 else ""))
                year = ds.get("year") or "—"
                description = (ds.get("description") or "").strip()
                if len(description) > 400:
                    description = description[:400].rstrip() + "…"
                description_html = html.escape(description) if description else "<em>No description available.</em>"
                doi = ds.get("doi")
                dataset_url = ds.get("dataset_url")
                formats = ds.get("file_formats") or []

                link_parts = []
                if dataset_url:
                    link_parts.append(f'<a href="{html.escape(dataset_url)}" target="_blank">View dataset</a>')
                if doi:
                    link_parts.append(f'<a href="https://doi.org/{html.escape(doi)}" target="_blank">DOI</a>')
                links_html = " · ".join(link_parts) if link_parts else ""

                source = html.escape(ds.get("source") or "")
                source_colors = {
                    "Zenodo": "#1457a8",
                    "DataCite": "#00a19a",
                    "OpenAlex Datasets": "#6a1b9a",
                }
                source_color = source_colors.get(ds.get("source", ""), "#555")
                source_badge = (
                    f'<span style="background:{source_color};color:#fff;font-size:0.7rem;'
                    f'font-weight:700;padding:0.2rem 0.55rem;border-radius:999px;margin-left:0.5rem;">'
                    f'{source}</span>'
                ) if source else ""
                format_badges = "".join(
                    f'<span class="dataset-format-chip">{html.escape(fmt)}</span>'
                    for fmt in formats
                )

                with st.expander(f"{i}. {ds.get('title', 'Untitled')} ({year})", expanded=(i == 1)):
                    st.markdown(
                        f"""
                        <div style="margin-bottom:0.5rem;">
                            {source_badge}
                            <span style="color:var(--text-secondary);font-size:0.85rem;margin-left:0.35rem;">{author_str}</span>
                            &nbsp;·&nbsp;
                            <span style="color:var(--text-secondary);font-size:0.85rem;">{year}</span>
                        </div>
                        <p style="font-size:0.9rem;line-height:1.6;color:var(--text-secondary);margin-bottom:0.6rem;">{description_html}</p>
                        <div style="margin-bottom:0.5rem;">{format_badges}</div>
                        <p style="font-size:0.82rem;">{links_html}</p>
                        """,
                        unsafe_allow_html=True,
                    )

        if st.button("Clear results", key="clear_dataset_search"):
            st.session_state.dataset_results = None
            st.session_state.dataset_query_used = ""
            st.session_state.dataset_summary = ""
            st.rerun()


# ── Tab 4: Rank Papers ───────────────────────────────────────────
elif selected_page == "Rank Papers":
    st.markdown("""
    <div class="qa-hero">
        <h3>Paper Relevance Ranker</h3>
        <p>Upload your collected PDFs and enter your research question.
        The AI ranks every paper by how well it fits your thesis and explains why.</p>
    </div>
    """, unsafe_allow_html=True)

    if "rank_results" not in st.session_state:
        st.session_state.rank_results = None
    if "rank_question" not in st.session_state:
        st.session_state.rank_question = ""

    with st.container(border=True):
        rank_question = st.text_area(
            "Your research question or thesis topic",
            placeholder="e.g. How does social media use affect mental health in adolescents?",
            height=90,
            label_visibility="collapsed",
        )
        rank_files = st.file_uploader(
            "Upload research PDFs (up to 20)",
            type=["pdf"],
            accept_multiple_files=True,
            key="rank_uploader",
        )
        rank_btn = st.button(
            "Rank Papers",
            type="primary",
            disabled=not (rank_question.strip() and rank_files),
            use_container_width=False,
        )

    if rank_btn and rank_question.strip() and rank_files:
        with st.spinner(f"Analysing {len(rank_files)} paper(s)… this may take a moment."):
            files_payload = [
                ("files", (f.name, f.getvalue(), "application/pdf"))
                for f in rank_files
            ]
            try:
                import requests as _req
                raw = _req.post(
                    f"{BACKEND_URL}/rank/papers",
                    data={"question": rank_question},
                    files=files_payload,
                    timeout=120,
                )
                raw.raise_for_status()
                result = raw.json()
            except Exception as e:
                st.error(f"Ranking failed: {e}")
                result = None

        if result:
            st.session_state.rank_results = result["papers"]
            st.session_state.rank_question = result["question"]

    if st.session_state.rank_results:
        papers = st.session_state.rank_results
        q = html.escape(st.session_state.rank_question)

        st.markdown(
            f'<p class="qa-toolbar-meta">Ranked {len(papers)} paper(s) for: <strong>{q}</strong></p>',
            unsafe_allow_html=True,
        )

        label_colors = {
            "Highly Relevant":    ("#e6f4ec", "#0a7c42", "#b8e6cc"),
            "Relevant":           ("#e8f0fb", "#1a56a0", "#b8d0f0"),
            "Somewhat Relevant":  ("#fef6e8", "#a05c00", "#f0d9a8"),
            "Less Relevant":      ("#fdecea", "#c0392b", "#f5c6c0"),
        }

        for rank, paper in enumerate(papers, 1):
            label = paper["label"]
            score_pct = min(100, max(0, int(paper["score"] * 100)))
            bg, fg, border = label_colors.get(label, ("#f0f4f8", "#333", "#ccc"))
            filename = html.escape(paper["filename"])
            explanation = html.escape(paper["explanation"])
            key_quote = html.escape(paper["key_quote"])

            with st.expander(f"{rank}. {paper['filename']}", expanded=(rank <= 3)):
                st.markdown(
                    f"""
                    <div style="display:flex;align-items:center;gap:0.75rem;margin-bottom:0.85rem;flex-wrap:wrap;">
                        <span style="background:{bg};color:{fg};border:1px solid {border};
                              font-size:0.78rem;font-weight:700;padding:0.25rem 0.75rem;
                              border-radius:999px;">{html.escape(label)}</span>
                        <div style="flex:1;min-width:120px;max-width:200px;height:8px;
                                    background:var(--border);border-radius:999px;overflow:hidden;">
                            <div style="width:{score_pct}%;height:100%;
                                        background:linear-gradient(90deg,var(--accent),var(--accent-hover));
                                        border-radius:999px;"></div>
                        </div>
                        <span style="font-size:0.82rem;color:var(--accent-hover);font-weight:600;">
                            {score_pct}% match
                        </span>
                    </div>
                    <p style="font-size:0.9rem;line-height:1.6;color:var(--text-secondary);margin-bottom:0.6rem;">
                        {explanation}
                    </p>
                    <div style="background:var(--bg-card-alt);border-left:3px solid var(--accent);
                                padding:0.6rem 0.9rem;border-radius:0 6px 6px 0;
                                font-size:0.85rem;color:var(--text-secondary);font-style:italic;">
                        "{key_quote}"
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

        st.markdown('<p class="stats-section-label">Export results</p>', unsafe_allow_html=True)
        exp_col1, exp_col2, exp_col3 = st.columns([1, 1, 2], vertical_alignment="center")
        with exp_col1:
            st.download_button(
                "Download Word",
                data=_rankings_to_word(st.session_state.rank_question, papers),
                file_name="paper_rankings.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with exp_col2:
            st.download_button(
                "Download PDF",
                data=_rankings_to_pdf(st.session_state.rank_question, papers),
                file_name="paper_rankings.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        with exp_col3:
            if st.button("Clear rankings", key="clear_rank", use_container_width=True):
                st.session_state.rank_results = None
                st.session_state.rank_question = ""
                st.rerun()


# ── Tab 5: Statistical Analysis ──────────────────────────────────
elif selected_page == "Statistical Analysis":
    if not st.session_state.csv_session_id:
        st.markdown(
            """
            <div class="stats-hero">
                <h3>Statistical Analysis</h3>
                <p>Ask a question in plain language — the AI selects and runs the appropriate test.</p>
            </div>
            <div class="stats-empty-state">
                <h4>No dataset loaded</h4>
                <p>Upload a CSV or XLSX file in the sidebar to run hypothesis tests.</p>
                <ol>
                    <li>Open the sidebar and upload your data file</li>
                    <li>Review columns in the <strong>Data Preview</strong> tab</li>
                    <li>Ask a question here or pick a suggested example</li>
                </ol>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        fname = html.escape(st.session_state.csv_filename)
        nrows = st.session_state.csv_rows
        ncols = len(st.session_state.csv_columns)
        st.markdown(
            f"""
            <div class="stats-hero">
                <h3>Statistical Analysis</h3>
                <p>Ask a question in plain language — the AI selects the test, runs it, and explains the results.</p>
                <span class="stats-dataset-pill">📊 {fname} · {nrows:,} rows · {ncols} columns</span>
            </div>
            """,
            unsafe_allow_html=True,
        )

        with st.container(border=True):
            if not st.session_state.example_questions:
                with st.spinner("Generating example questions for your dataset…"):
                    st.session_state.example_questions = generate_example_questions(
                        st.session_state.csv_session_id,
                        st.session_state.csv_columns,
                        st.session_state.csv_dtypes,
                    )

            st.markdown('<p class="stats-section-label">Suggested questions</p>', unsafe_allow_html=True)
            with st.expander("Browse examples for your dataset", expanded=False):
                st.markdown('<div class="stats-example-row">', unsafe_allow_html=True)
                for row_start in range(0, len(st.session_state.example_questions), 2):
                    ex_cols = st.columns(2, gap="small", vertical_alignment="top")
                    for col_idx, ex in enumerate(
                        st.session_state.example_questions[row_start : row_start + 2]
                    ):
                        with ex_cols[col_idx]:
                            if st.button(
                                ex,
                                key=f"ex_{row_start + col_idx}",
                                use_container_width=True,
                                type="secondary",
                            ):
                                st.session_state.stat_question = ex
                                st.rerun()
                st.markdown("</div>", unsafe_allow_html=True)

            st.markdown('<p class="stats-section-label">Your analysis</p>', unsafe_allow_html=True)
            st.markdown(
                '<p class="stats-input-hint">Describe what you want to test — comparisons, correlations, '
                "regression, or associations. Use exact column names when possible.</p>",
                unsafe_allow_html=True,
            )
            stat_question = st.text_area(
                "Research question",
                key="stat_question",
                placeholder="e.g. Does smoking predict disease?",
                height=88,
                label_visibility="collapsed",
            )

            alpha_col, btn_col = st.columns([2, 1], gap="medium", vertical_alignment="bottom")
            with alpha_col:
                alpha = st.selectbox("Significance level (α)", [0.05, 0.01, 0.10], index=0)
            with btn_col:
                run_btn = st.button(
                    "Run Analysis",
                    type="primary",
                    disabled=not stat_question.strip(),
                    use_container_width=True,
                )

        if run_btn and stat_question.strip():
            with st.spinner("Selecting test · Running analysis · Generating explanation…"):
                result = backend(
                    "POST", "/stats/analyze",
                    json={"session_id": st.session_state.csv_session_id, "question": stat_question},
                )

            if result:
                st.session_state.last_stats_result = result
                if not (result.get("test_name") == "unsupported" or result.get("error")):
                    st.session_state.stats_history.append({
                        "question": stat_question,
                        "result": result,
                        "session_id": st.session_state.csv_session_id,
                    })

        # Render the most recent result unconditionally (not gated on run_btn) so it
        # survives reruns triggered by unrelated widgets, e.g. the chart Q&A "Ask" button.
        last_result = st.session_state.last_stats_result
        if last_result:
            if last_result.get("test_name") == "unsupported":
                msg = html.escape(last_result.get("plain_explanation") or last_result.get("error", ""))
                st.markdown(
                    f'<div class="stats-unsupported-box"><strong>Not supported</strong><br>{msg}</div>',
                    unsafe_allow_html=True,
                )
                rat = last_result.get("rationale")
                if rat and rat != last_result.get("plain_explanation"):
                    st.caption(rat)
            elif last_result.get("error"):
                err = html.escape(str(last_result["error"]))
                st.markdown(
                    f'<div class="stats-error-box"><strong>Analysis could not be completed</strong><br>{err}</div>',
                    unsafe_allow_html=True,
                )
            else:
                render_stats_result(last_result, session_id=st.session_state.csv_session_id)

        if st.session_state.stats_history:
            st.markdown('<p class="stats-section-label">Session history</p>', unsafe_allow_html=True)
            hist_header, hist_exp_csv, hist_exp_xlsx, hist_clear = st.columns([3, 1, 1, 1], vertical_alignment="center")
            with hist_exp_csv:
                st.download_button(
                    "Export CSV",
                    data=_stats_to_bytes(st.session_state.stats_history, "csv"),
                    file_name="analysis_results.csv",
                    mime="text/csv",
                    use_container_width=True,
                )
            with hist_exp_xlsx:
                st.download_button(
                    "Export XLSX",
                    data=_stats_to_bytes(st.session_state.stats_history, "xlsx"),
                    file_name="analysis_results.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
            with hist_clear:
                if st.button("Clear history", key="clear_stats", use_container_width=True):
                    st.session_state.stats_history = []
                    st.rerun()

            for h in reversed(st.session_state.stats_history[-5:]):
                r = h["result"]
                q_short = html.escape(h["question"][:90] + ("…" if len(h["question"]) > 90 else ""))
                test_lbl = html.escape(str(r.get("test_display_name", "—")))
                p_disp = r.get("p_value")
                p_str = f"{p_disp:.4f}" if p_disp is not None else "N/A"
                with st.expander(f"{test_lbl} · p = {p_str}"):
                    st.markdown(
                        f'<p class="stats-history-item"><strong>Question:</strong> {q_short}</p>',
                        unsafe_allow_html=True,
                    )
                    if r.get("error") or r.get("test_name") == "unsupported":
                        st.markdown(r.get("plain_explanation") or r.get("error", ""))
                    elif r.get("plain_explanation"):
                        st.markdown(r["plain_explanation"])
                    else:
                        render_stats_result(r, session_id=h.get("session_id"))


# ── Tab 6: Statistics Handbook ───────────────────────────────────
elif selected_page == "Statistics Handbook":
    st.markdown(
        '<h2 style="margin:0 0 0.2rem;font-size:1.5rem;color:var(--text-primary);">Stats handbook</h2>'
        '<p style="margin:0 0 1.25rem;color:var(--text-secondary);font-size:0.92rem;">'
        'Look up core concepts in plain language, or ask about a specific term.</p>',
        unsafe_allow_html=True,
    )

    last_result = st.session_state.last_stats_result
    groundable = bool(
        last_result
        and last_result.get("test_name")
        and last_result.get("test_name") != "unsupported"
        and not last_result.get("error")
    )

    def _handbook_payload(question: str, ground: bool) -> dict:
        payload = {"question": question, "history": st.session_state.handbook_chat_history[:-1]}
        if ground:
            payload.update({
                "test_name": last_result.get("test_name"),
                "variables_used": last_result.get("variables_used") or {},
                "rationale": last_result.get("rationale"),
                "statistic": last_result.get("statistic"),
                "p_value": last_result.get("p_value"),
                "alpha": last_result.get("alpha", 0.05),
                "significant": last_result.get("significant"),
                "assumption_checks": last_result.get("assumption_checks") or [],
            })
        return payload

    def _ask_handbook(question: str, ground: bool) -> None:
        st.session_state.handbook_chat_history.append({"role": "user", "content": question})
        with st.spinner("Thinking…"):
            result = backend("POST", "/handbook/ask", json=_handbook_payload(question, ground))
        if result:
            st.session_state.handbook_chat_history.append({"role": "assistant", "content": result["answer"]})

    with st.container(border=True):
        st.markdown(
            '<div class="handbook-ask-title">❓ Ask about a concept</div>'
            '<p class="handbook-ask-subtitle">e.g. "What\'s the difference between a t-test and ANOVA?"</p>',
            unsafe_allow_html=True,
        )

        ask_col, btn_col = st.columns([5, 1], vertical_alignment="bottom")
        with ask_col:
            concept_question = st.text_input(
                "Ask about a concept",
                placeholder="What is a p-value?",
                key="handbook_concept_input",
                label_visibility="collapsed",
            )
        with btn_col:
            ask_clicked = st.button("➤ Ask", key="handbook_ask_btn", type="primary", use_container_width=True)

        ground_in_result = False
        if groundable:
            ground_in_result = st.checkbox(
                f"Ground answers in my last result ({last_result.get('test_display_name', '')})",
                value=True,
                key="handbook_ground_checkbox",
            )

        if ask_clicked and concept_question.strip():
            _ask_handbook(concept_question.strip(), ground_in_result)
            st.rerun()

        if not st.session_state.handbook_chat_history:
            starter_questions = [
                "What is a p-value?",
                "What's the difference between a t-test and ANOVA?",
                "What does 'statistically significant' actually mean?",
            ]
            pill_cols = st.columns(len(starter_questions))
            for idx, sq in enumerate(starter_questions):
                with pill_cols[idx]:
                    if st.button(sq, key=f"handbook_starter_{idx}", type="secondary"):
                        _ask_handbook(sq, ground_in_result)
                        st.rerun()
        else:
            last_answer = next(
                (m["content"] for m in reversed(st.session_state.handbook_chat_history) if m["role"] == "assistant"),
                None,
            )
            if last_answer:
                st.markdown(f'<div class="handbook-answer-box">{last_answer}</div>', unsafe_allow_html=True)
            if st.button("Clear chat", key="clear_handbook_chat"):
                st.session_state.handbook_chat_history = []
                st.rerun()

    if st.session_state.handbook_concepts is None:
        with st.spinner("Loading handbook…"):
            resp = backend("GET", "/handbook/concepts")
        st.session_state.handbook_concepts = (resp or {}).get("concepts", [])

    concepts = st.session_state.handbook_concepts or []
    core_concepts = [c for c in concepts if c.get("category") == "concept"]
    test_concepts = [c for c in concepts if c.get("category") == "test"]

    def _concept_grid(items: list, icon: str) -> None:
        cols = st.columns(2, gap="small")
        for idx, c in enumerate(items):
            with cols[idx % 2]:
                st.markdown(
                    f"""
                    <div class="concept-card">
                        <div class="concept-card-title">
                            <span class="concept-card-icon">{icon}</span>{html.escape(c['title'])}
                        </div>
                        <div class="concept-card-body">{html.escape(c['what'])}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                with st.expander("More", expanded=False):
                    label = "When to use it" if c.get("category") == "test" else "Why it matters"
                    st.markdown(f"**{label}:** {c['why_it_matters']}")
                    st.markdown(f"**Example:** {c['example']}")

    st.markdown('<p class="stats-section-label">Core concepts</p>', unsafe_allow_html=True)
    _concept_grid(core_concepts, "📘")

    st.markdown('<p class="stats-section-label">Hypothesis tests</p>', unsafe_allow_html=True)
    _concept_grid(test_concepts, "🧪")


# ── Tab 7: Data Preview ──────────────────────────────────────────
elif selected_page == "Data Preview":
    st.markdown("### Dataset preview")
    if not st.session_state.csv_session_id:
        st.info("Upload a CSV or XLSX file to see a preview here.")
    else:
        st.markdown(
            f"**{st.session_state.csv_filename}** · "
            f"{st.session_state.csv_rows} rows · {len(st.session_state.csv_columns)} columns"
        )
        if st.session_state.csv_preview:
            st.dataframe(pd.DataFrame(st.session_state.csv_preview), use_container_width=True)

        st.markdown("**Column types:**")
        dtype_df = pd.DataFrame(
            [{"Column": col, "Type": st.session_state.csv_dtypes.get(col, "")}
             for col in st.session_state.csv_columns]
        )
        st.dataframe(dtype_df, use_container_width=True, hide_index=True)
